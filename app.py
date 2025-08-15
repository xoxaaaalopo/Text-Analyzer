from flask import Flask, request, jsonify
from flask_cors import CORS
import re
from collections import Counter
import zipfile
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import tempfile
import os
import json
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet
from nltk import pos_tag
import spacy
import csv
import pdfplumber
from docx import Document
from striprtf.striprtf import rtf_to_text

app = Flask(__name__)
CORS(app)

with open('advanced_words.csv', 'r', encoding='utf-8') as f:
    reader = csv.reader(f)
    advanced_words_set = set(row[0] for row in reader)

nlp = spacy.load("en_core_web_lg")

with open('Frequency_FLT.json', 'r', encoding='utf-8') as f:
    freq_data = json.load(f)

freq_dict = {entry[0]: entry[2]['frequency'] for entry in freq_data}

def lemmatize_words(text: str) -> list[str]:
    doc = nlp(text)
    return [token.lemma_ for token in doc if token.is_alpha]

class VocabularyAnalyzer:
    def extract_text_from_pdf(self, file_content):
        """Extract text from PDF using pdfplumber"""
        with tempfile.NamedTemporaryFile() as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            
            with pdfplumber.open(temp_file.name) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return ' '.join(text_parts)
    
    def extract_text_from_docx(self, file_content):
        """Extract text from DOCX file"""
        with tempfile.NamedTemporaryFile() as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            
            doc = Document(temp_file.name)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            return ' '.join(text_parts)
    
    def extract_text_from_rtf(self, file_content):
        """Extract text from RTF file"""
        rtf_text = file_content.decode('utf-8', errors='ignore')
        return rtf_to_text(rtf_text)

    def extract_text_from_srt(self, file_content):
        """Extract text from SRT subtitle file"""
        text = file_content.decode('utf-8', errors='ignore')

        text = re.sub(r'\d+\n\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3}\n', '', text)

        text = re.sub(r'<[^>]+>', '', text)
        
        text = re.sub(r'\n+', ' ', text)
        
        return text.strip()

    def extract_text_from_epub(self, file_content):
        """Extract text from EPUB file"""
        extracted_text = []
        
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            with zipfile.ZipFile(temp_file_path, 'r') as epub_zip:
                # Find all HTML/XHTML files in the EPUB
                for file_info in epub_zip.infolist():
                    if file_info.filename.endswith(('.html', '.xhtml', '.htm')):
                        with epub_zip.open(file_info.filename) as html_file:
                            html_content = html_file.read().decode('utf-8', errors='ignore')
                            # Parse HTML and extract text
                            soup = BeautifulSoup(html_content, 'html.parser')
                            text = soup.get_text()
                            extracted_text.append(text)
        finally:
            os.unlink(temp_file_path)
        
        return ' '.join(extracted_text)
    
    def extract_text_by_format(self, file_content, file_extension):
        """Route to appropriate extraction method based on file format"""
        extractors = {
            'pdf': self.extract_text_from_pdf,
            'docx': self.extract_text_from_docx,
            'rtf': self.extract_text_from_rtf,
            'srt': self.extract_text_from_srt,
            'epub': self.extract_text_from_epub,
            'txt': self.extract_text_from_txt,
        }
        
        extractor = extractors.get(file_extension)
        if not extractor:
            raise ValueError(f'Unsupported file format: {file_extension}')
        
        return extractor(file_content)

    def extract_text_from_txt(self, file_content):
        """Extract text from plain text file"""
        return file_content.decode('utf-8', errors='ignore')

    def clean_and_tokenize(self, text):
        """Clean text and extract words"""
        text = text.lower()
        
        text = re.sub(r'[^a-zA-Z\s]', ' ', text)
        
        words = [word for word in text.split()]

        return words

    def analyze_vocabulary(self, text):
        words = self.clean_and_tokenize(text)

        unique_words = list(set(lemmatize_words(text)))

        advanced_words = [word for word in unique_words if word in advanced_words_set]

        word_data = []
        for word in unique_words:
            freq = freq_dict.get(word)
            if freq is not None:
                word_data.append({
                    'word': word,
                    'frequency': freq
                })
        
        word_data.sort(key=lambda x: x['frequency'] if isinstance(x['frequency'], int) else -1, reverse=True)

        return {
            'total_num': len(words),
            'unique_num': len(unique_words),
            'sorted_words': word_data,
            'advanced_words': advanced_words
        }

@app.route('/api/analyzeFile', methods=['POST'])
def analyze_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        file_extension = file.filename.lower().split('.')[-1]
        file_content = file.read()
        
        analyzer = VocabularyAnalyzer()

        text = analyzer.extract_text_by_format(file_content, file_extension)
        
        results = analyzer.analyze_vocabulary(text)
        
        return jsonify(results)
    
    except Exception as e:
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/api/analyzeText', methods=['POST'])
def analyze_text():
    data = request.json
    text = data.get('text')
    analyzer = VocabularyAnalyzer()

    results = analyzer.analyze_vocabulary(text)

    return jsonify(results)

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy'})

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5002)

# Requirements for requirements.txt:
"""
Flask==2.3.3
Flask-CORS==4.0.0
beautifulsoup4==4.12.2
lxml==4.9.3
"""